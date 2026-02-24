"""
Markdown to DOCX converter for TSC-001_Audio_Input_Qualification_Layer_v04.md

Converts the markdown technical safety concept document to a Word .docx file
with proper formatting for headings, tables, code blocks, lists, and inline styles.

Requirements: python-docx >= 1.1.0
"""

import os
import re
from docx import Document
from docx.shared import Inches, Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def set_narrow_margins(doc):
    """Set narrow margins (0.5 inch) on all sections."""
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)


def add_inline_formatting(paragraph, text):
    """Parse inline markdown formatting and add runs to a paragraph.

    Handles: ***bold-italic***, **bold**, *italic*, `code`, and plain text.
    Processes mixed formatting within a single line.
    """
    # Pattern to match inline formatting elements
    # Order matters: bold-italic first, then bold, then italic, then code
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'   # ***bold-italic***
        r'|(\*\*(.+?)\*\*)'       # **bold**
        r'|(\*(.+?)\*)'           # *italic*
        r'|(`(.+?)`)'             # `code`
    )

    last_end = 0
    for match in pattern.finditer(text):
        # Add plain text before this match
        if match.start() > last_end:
            plain = text[last_end:match.start()]
            if plain:
                paragraph.add_run(plain)

        if match.group(2) is not None:
            # ***bold-italic***
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
        elif match.group(4) is not None:
            # **bold**
            run = paragraph.add_run(match.group(4))
            run.bold = True
        elif match.group(6) is not None:
            # *italic*
            run = paragraph.add_run(match.group(6))
            run.italic = True
        elif match.group(8) is not None:
            # `code`
            run = paragraph.add_run(match.group(8))
            run.font.name = 'Consolas'
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)

        last_end = match.end()

    # Add remaining plain text
    if last_end < len(text):
        remaining = text[last_end:]
        if remaining:
            paragraph.add_run(remaining)

    # If no formatting was found at all and text is empty, do nothing
    # If no formatting was found but there is text, add it as plain
    if last_end == 0 and text:
        # Clear runs (shouldn't have any yet from this call path, but be safe)
        # Actually the paragraph might already have the run added above
        pass


def add_cell_formatting(cell, text, bold_all=False, font_size=Pt(8)):
    """Add formatted text to a table cell, handling inline markdown."""
    cell.text = ""  # Clear default text
    paragraph = cell.paragraphs[0]

    # Set cell font size via paragraph format
    if bold_all:
        # For header cells, apply bold to all text but still parse inline formatting
        # Simple approach: add bold runs
        run = paragraph.add_run(text.strip())
        run.bold = True
        run.font.size = font_size
        run.font.name = 'Calibri'
    else:
        # Parse inline formatting
        add_inline_formatting_cell(paragraph, text.strip(), font_size)


def add_inline_formatting_cell(paragraph, text, font_size=Pt(8)):
    """Parse inline markdown formatting for table cells with specific font size."""
    pattern = re.compile(
        r'(\*\*\*(.+?)\*\*\*)'
        r'|(\*\*(.+?)\*\*)'
        r'|(\*(.+?)\*)'
        r'|(`(.+?)`)'
    )

    last_end = 0
    for match in pattern.finditer(text):
        if match.start() > last_end:
            plain = text[last_end:match.start()]
            if plain:
                run = paragraph.add_run(plain)
                run.font.size = font_size
                run.font.name = 'Calibri'

        if match.group(2) is not None:
            run = paragraph.add_run(match.group(2))
            run.bold = True
            run.italic = True
            run.font.size = font_size
            run.font.name = 'Calibri'
        elif match.group(4) is not None:
            run = paragraph.add_run(match.group(4))
            run.bold = True
            run.font.size = font_size
            run.font.name = 'Calibri'
        elif match.group(6) is not None:
            run = paragraph.add_run(match.group(6))
            run.italic = True
            run.font.size = font_size
            run.font.name = 'Calibri'
        elif match.group(8) is not None:
            run = paragraph.add_run(match.group(8))
            run.font.name = 'Consolas'
            run.font.size = Pt(7)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)

        last_end = match.end()

    if last_end < len(text):
        remaining = text[last_end:]
        if remaining:
            run = paragraph.add_run(remaining)
            run.font.size = font_size
            run.font.name = 'Calibri'

    if last_end == 0 and text:
        pass  # Already handled by the remaining text block above


def parse_table_row(line):
    """Parse a markdown table row into a list of cell contents."""
    # Strip leading/trailing pipes and split
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    cells = [c.strip() for c in line.split('|')]
    return cells


def is_separator_row(line):
    """Check if a table row is a separator (e.g., |---|---|)."""
    stripped = line.strip()
    # Remove pipes and check if remaining is only dashes, colons, spaces
    content = stripped.replace('|', '').replace('-', '').replace(':', '').replace(' ', '')
    return len(content) == 0 and '-' in stripped


def add_table(doc, rows):
    """Add a formatted table to the document.

    rows: list of lists of cell text. First row is header.
    """
    if not rows or not rows[0]:
        return

    num_cols = len(rows[0])
    num_rows = len(rows)

    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Calculate available width (page width minus margins)
    section = doc.sections[-1]
    available_width = section.page_width - section.left_margin - section.right_margin

    # Set column widths proportionally based on content
    # First, estimate content widths
    col_widths = []
    for col_idx in range(num_cols):
        max_len = 0
        for row in rows:
            if col_idx < len(row):
                max_len = max(max_len, len(row[col_idx]))
        col_widths.append(max(max_len, 5))  # Minimum width

    total_content = sum(col_widths)
    for col_idx in range(num_cols):
        width = int(available_width * col_widths[col_idx] / total_content)
        for row_idx in range(num_rows):
            table.cell(row_idx, col_idx).width = width

    # Populate cells
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.cell(row_idx, col_idx)
                is_header = (row_idx == 0)
                add_cell_formatting(cell, cell_text, bold_all=is_header, font_size=Pt(8))

                # Add shading to header row
                if is_header:
                    shading = parse_xml(
                        f'<w:shd {nsdecls("w")} w:fill="D9E2F3" w:val="clear"/>'
                    )
                    cell._element.get_or_add_tcPr().append(shading)

    # Add a small paragraph after the table for spacing
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)


def add_code_block(doc, code_lines):
    """Add a code block with monospace formatting and preserved whitespace."""
    for line in code_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = Pt(11)

        # Set paragraph shading (light gray background)
        pPr = p._element.get_or_add_pPr()
        shading = parse_xml(
            f'<w:shd {nsdecls("w")} w:fill="F2F2F2" w:val="clear"/>'
        )
        pPr.append(shading)

        # Use non-breaking spaces to preserve leading whitespace
        # But first, just add the line content preserving spaces
        run = p.add_run(line)
        run.font.name = 'Consolas'
        run.font.size = Pt(7.5)
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def add_horizontal_rule(doc):
    """Add a horizontal rule (thin line) to the document."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Add a bottom border to simulate HR
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="999999"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def get_list_indent(line):
    """Get the indentation level and content of a list item."""
    # Count leading spaces
    stripped = line.lstrip()
    indent = len(line) - len(stripped)
    level = indent // 2  # Each level is ~2 spaces (or 4, adjust as needed)
    if indent >= 4:
        level = indent // 4
    elif indent >= 2:
        level = 1
    else:
        level = 0
    return level


def convert_md_to_docx(md_path, docx_path):
    """Main conversion function."""
    # Read the markdown file
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    set_narrow_margins(doc)

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10)

    # Configure heading styles
    for level in range(1, 5):
        heading_style = doc.styles[f'Heading {level}']
        heading_style.font.name = 'Calibri'
        if level == 1:
            heading_style.font.size = Pt(20)
            heading_style.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        elif level == 2:
            heading_style.font.size = Pt(16)
            heading_style.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
        elif level == 3:
            heading_style.font.size = Pt(13)
            heading_style.font.color.rgb = RGBColor(0x2E, 0x4A, 0x7A)
        elif level == 4:
            heading_style.font.size = Pt(11)
            heading_style.font.color.rgb = RGBColor(0x2E, 0x4A, 0x7A)

    i = 0
    first_h2 = True  # Track first H2 to avoid page break before first content

    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Code blocks (triple backtick)
        if line.strip().startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            i += 1  # Skip closing ```
            add_code_block(doc, code_lines)
            continue

        # Horizontal rule
        if line.strip() == '---' or line.strip() == '***' or line.strip() == '___':
            add_horizontal_rule(doc)
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,4})\s+(.+)$', line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()

            # Add page break before H2 headings (except the very first one)
            if level == 2:
                if first_h2:
                    first_h2 = False
                else:
                    # Add page break
                    p = doc.add_paragraph()
                    run = p.add_run()
                    run.add_break(docx.enum.text.WD_BREAK.PAGE)

            heading = doc.add_heading(text, level=level)
            # Parse inline formatting in headings
            # Clear default run and re-add with formatting
            if '**' in text or '*' in text or '`' in text:
                # Remove all runs
                for run in heading.runs:
                    run.text = ''
                add_inline_formatting(heading, text)
                # Reapply heading font to all runs
                for run in heading.runs:
                    run.font.name = 'Calibri'
                    if level == 1:
                        run.font.size = Pt(20)
                        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
                    elif level == 2:
                        run.font.size = Pt(16)
                        run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
                    elif level == 3:
                        run.font.size = Pt(13)
                        run.font.color.rgb = RGBColor(0x2E, 0x4A, 0x7A)
                    elif level == 4:
                        run.font.size = Pt(11)
                        run.font.color.rgb = RGBColor(0x2E, 0x4A, 0x7A)

            i += 1
            continue

        # Tables
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                if is_separator_row(row_line):
                    i += 1
                    continue
                cells = parse_table_row(row_line)
                table_rows.append(cells)
                i += 1
            if table_rows:
                add_table(doc, table_rows)
            continue

        # Ordered lists
        ol_match = re.match(r'^(\s*)\d+\.\s+(.+)$', line)
        if ol_match:
            # Collect all consecutive ordered list items
            while i < len(lines):
                ol_line = lines[i].rstrip('\n')
                ol_item_match = re.match(r'^(\s*)\d+\.\s+(.+)$', ol_line)
                if ol_item_match:
                    indent = len(ol_item_match.group(1))
                    content = ol_item_match.group(2)
                    level = indent // 3 if indent >= 3 else 0

                    p = doc.add_paragraph(style='List Number')
                    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after = Pt(1)

                    # Clear default text and add formatted text
                    p.text = ''
                    add_inline_formatting(p, content)

                    i += 1
                elif ol_line.strip() == '':
                    i += 1
                    break
                else:
                    break
            continue

        # Unordered lists
        ul_match = re.match(r'^(\s*)[-*]\s+(.+)$', line)
        if ul_match:
            while i < len(lines):
                ul_line = lines[i].rstrip('\n')
                ul_item_match = re.match(r'^(\s*)[-*]\s+(.+)$', ul_line)
                if ul_item_match:
                    indent = len(ul_item_match.group(1))
                    content = ul_item_match.group(2)
                    level = 0
                    if indent >= 4:
                        level = indent // 4
                    elif indent >= 2:
                        level = 1

                    p = doc.add_paragraph(style='List Bullet')
                    p.paragraph_format.left_indent = Inches(0.25 + level * 0.25)
                    p.paragraph_format.space_before = Pt(1)
                    p.paragraph_format.space_after = Pt(1)

                    # Clear default text and add formatted text
                    p.text = ''
                    add_inline_formatting(p, content)

                    i += 1
                elif ul_line.strip() == '':
                    i += 1
                    break
                else:
                    break
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        add_inline_formatting(p, line.strip())

        i += 1

    # Save
    doc.save(docx_path)
    print(f"Successfully converted to: {docx_path}")


if __name__ == '__main__':
    import docx.enum.text

    # Paths relative to script location
    script_dir = os.path.dirname(os.path.abspath(__file__))
    md_file = os.path.join(script_dir, 'TSC-001_Audio_Input_Qualification_Layer_v04.md')
    docx_file = os.path.join(script_dir, 'TSC-001_Audio_Input_Qualification_Layer_v05.docx')

    if not os.path.exists(md_file):
        print(f"ERROR: Markdown file not found: {md_file}")
        exit(1)

    convert_md_to_docx(md_file, docx_file)
